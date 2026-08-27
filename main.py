"""
ACCO AI — Aqlli Buxgalter Yordamchisi (Telegram Bot)
----------------------------------------------------
Foydalanuvchi yuborgan PDF / Excel / Word hujjatlarini tahlil qilib,
Gemini AI yordamida o'zbek tilida professional audit xulosasini qaytaradi.

Sirlar (Replit Secrets):
  - TELEGRAM_BOT_TOKEN
  - GEMINI_API_KEY
"""

import io
import logging
import os

import pymupdf  # PyMuPDF
import google.generativeai as genai
from docx import Document as DocxDocument
from openpyxl import load_workbook
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import (
    ApplicationBuilder,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
)

# ---------------------------------------------------------------------------
# Sozlamalar
# ---------------------------------------------------------------------------
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
CHANNEL_USERNAME = "@AccoAI"

LANGUAGES = {
    "uz": "O'zbekcha",
    "ru": "Русский",
    "en": "English",
}

TEXTS = {
    "choose_language": {
        "uz": "Tilni tanlang:",
        "ru": "Выберите язык:",
        "en": "Choose your language:",
    },
    "must_subscribe": {
        "uz": "ACCO AI botidan foydalanish uchun @AccoAI kanaliga majburiy obuna bo'ling.\nObuna bo'lgach, «Obunani tekshirish» tugmasini bosing.",
        "ru": "Для использования ACCO AI необходимо подписаться на канал @AccoAI.\nПосле подписки нажмите «Проверить подписку».",
        "en": "To use ACCO AI, you must subscribe to the @AccoAI channel.\nAfter subscribing, press “Check subscription”.",
    },
    "join_channel": {"uz": "Kanalga obuna bo'lish", "ru": "Подписаться на канал", "en": "Subscribe to channel"},
    "check_subscription": {"uz": "Obunani tekshirish", "ru": "Проверить подписку", "en": "Check subscription"},
    "subscription_ok": {
        "uz": "Obuna tasdiqlandi. Endi hujjat yuborishingiz mumkin.",
        "ru": "Подписка подтверждена. Теперь отправьте документ.",
        "en": "Subscription confirmed. You can now send a document.",
    },
    "not_subscribed": {
        "uz": "Avval @AccoAI kanaliga obuna bo'ling.",
        "ru": "Сначала подпишитесь на канал @AccoAI.",
        "en": "Please subscribe to the @AccoAI channel first.",
    },
    "welcome": {
        "uz": "Salom! Men ACCO AI — aqlli buxgalter yordamchisiman.\n\nMenga PDF, Word, Excel hujjat yoki rasm yuboring.\n\nIkki hujjatni solishtirish uchun /compare buyrug'ini ishlating.",
        "ru": "Здравствуйте! Я ACCO AI — умный бухгалтерский помощник.\n\nОтправьте мне PDF, Word, Excel-документ или изображение.\n\nДля сравнения двух документов используйте команду /compare.",
        "en": "Hello! I am ACCO AI — your intelligent accounting assistant.\n\nSend me a PDF, Word, Excel document, or image.\n\nUse /compare to compare two documents.",
    },
    "choose_result": {
        "uz": "Tahlil tayyor. Natijani qanday ko'rishni xohlaysiz?",
        "ru": "Анализ готов. Как вы хотите получить результат?",
        "en": "The analysis is ready. How would you like to receive the result?",
    },
    "as_text": {"uz": "Matn ko'rinishida", "ru": "В виде текста", "en": "As text"},
    "as_file": {"uz": "Fayl ko'rinishida", "ru": "В виде файла", "en": "As a file"},
    "result_sending_text": {"uz": "Natija matn ko'rinishida yuborilmoqda...", "ru": "Результат отправляется в виде текста...", "en": "Sending the result as text..."},
    "result_sending_file": {"uz": "Natija fayl ko'rinishida yuborilmoqda...", "ru": "Результат отправляется в виде файла...", "en": "Sending the result as a file..."},
    "result_missing": {"uz": "Natija topilmadi. Iltimos, hujjatni qaytadan yuboring.", "ru": "Результат не найден. Отправьте документ ещё раз.", "en": "Result not found. Please send the document again."},
    "compare_started": {
        "uz": "Solishtirish rejimi yoqildi.\n\nIltimos, BIRINCHI hujjatni yuboring (PDF, Word, Excel yoki rasm).\nBekor qilish uchun: /cancel",
        "ru": "Режим сравнения включён.\n\nОтправьте ПЕРВЫЙ документ (PDF, Word, Excel или изображение).\nДля отмены: /cancel",
        "en": "Comparison mode is on.\n\nSend the FIRST document (PDF, Word, Excel, or image).\nTo cancel: /cancel",
    },
    "cancelled": {"uz": "Solishtirish rejimi bekor qilindi.", "ru": "Режим сравнения отменён.", "en": "Comparison mode cancelled."},
    "nothing_to_cancel": {"uz": "Bekor qilinadigan amal yo'q.", "ru": "Нет активного действия для отмены.", "en": "There is nothing to cancel."},
    "file_received": {"uz": "Fayl qabul qilindi, tayyorlanmoqda...", "ru": "Файл получен, подготовка...", "en": "File received, preparing it..."},
    "analysis_wait": {"uz": "Fayl qabul qilindi. AI tahlil qilmoqda, iltimos kuting...", "ru": "Файл получен. ИИ выполняет анализ, подождите...", "en": "File received. AI is analyzing it, please wait..."},
    "image_wait": {"uz": "Rasm qabul qilindi. AI vision orqali o'qib tahlil qilmoqda, iltimos kuting...", "ru": "Изображение получено. ИИ распознаёт и анализирует его, подождите...", "en": "Image received. AI is reading and analyzing it, please wait..."},
    "unsupported": {
        "uz": "Bu fayl turi qo'llab-quvvatlanmaydi.\nIltimos, PDF, Word (.docx), Excel (.xlsx) yoki rasm yuboring.",
        "ru": "Этот тип файла не поддерживается.\nОтправьте PDF, Word (.docx), Excel (.xlsx) или изображение.",
        "en": "This file type is not supported.\nPlease send a PDF, Word (.docx), Excel (.xlsx), or image.",
    },
    "no_data": {"uz": "Fayldan ma'lumot topib bo'lmadi. Iltimos, boshqa fayl yuboring.", "ru": "Не удалось найти данные в файле. Отправьте другой файл.", "en": "No data was found in the file. Please send another file."},
    "second_document": {
        "uz": "Birinchi hujjat qabul qilindi ({doc_type}).\n\nEndi IKKINCHI hujjatni yuboring.\nBekor qilish uchun: /cancel",
        "ru": "Первый документ получен ({doc_type}).\n\nТеперь отправьте ВТОРОЙ документ.\nДля отмены: /cancel",
        "en": "The first document was received ({doc_type}).\n\nNow send the SECOND document.\nTo cancel: /cancel",
    },
    "both_documents": {"uz": "Ikkala hujjat qabul qilindi. AI ularni solishtirmoqda, iltimos kuting...", "ru": "Оба документа получены. ИИ сравнивает их, подождите...", "en": "Both documents were received. AI is comparing them, please wait..."},
    "no_ai_response": {"uz": "AI hech qanday javob qaytarmadi. Iltimos, qaytadan urinib ko'ring.", "ru": "ИИ не вернул ответ. Попробуйте ещё раз.", "en": "AI returned no response. Please try again."},
    "ocr_wait": {"uz": "Hujjatda matn topilmadi — skaner qilingan PDF kabi ko'rinmoqda. AI vision orqali rasmlarni o'qimoqda, iltimos kuting...", "ru": "Текст не найден — похоже на сканированный PDF. ИИ распознаёт страницы, подождите...", "en": "No text found — this looks like a scanned PDF. AI is reading the pages, please wait..."},
    "empty_document": {"uz": "Hujjatdan matn topilmadi. Hujjat bo'sh bo'lishi mumkin — iltimos, boshqa fayl yuboring.", "ru": "Текст в документе не найден. Возможно, он пуст — отправьте другой файл.", "en": "No text was found in the document. It may be empty — please send another file."},
    "other_prompt": {"uz": "Iltimos, PDF, Word (.docx), Excel (.xlsx) hujjatini yoki rasm yuboring.", "ru": "Отправьте PDF, Word (.docx), Excel (.xlsx) документ или изображение.", "en": "Please send a PDF, Word (.docx), Excel (.xlsx) document, or image."},
}


def tr(context: ContextTypes.DEFAULT_TYPE, key: str) -> str:
    language = context.user_data.get("language", "uz")
    return TEXTS.get(key, {}).get(language, TEXTS.get(key, {}).get("uz", key))

if not TELEGRAM_BOT_TOKEN:
    raise RuntimeError("TELEGRAM_BOT_TOKEN env var (Replit Secret) topilmadi.")
if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY env var (Replit Secret) topilmadi.")

genai.configure(api_key=GEMINI_API_KEY)
PRIMARY_MODEL_NAME = "gemini-2.5-flash"
FALLBACK_MODEL_NAMES = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite"]
model = genai.GenerativeModel(PRIMARY_MODEL_NAME)


def _is_quota_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return (
        "429" in msg
        or "quota" in msg
        or "rate limit" in msg
        or "exhaust" in msg
        or "resourceexhausted" in msg
    )


def generate_with_fallback(parts):
    """Asosiy modelni chaqiradi; 429/quota xatolarida zaxira modellarga o'tadi."""
    try:
        return model.generate_content(parts)
    except Exception as exc:
        if not _is_quota_error(exc):
            raise
        last_exc = exc
        for name in FALLBACK_MODEL_NAMES:
            try:
                logger.warning(
                    "Asosiy model band (%s). Zaxira: %s",
                    PRIMARY_MODEL_NAME, name,
                )
                fb = genai.GenerativeModel(name)
                return fb.generate_content(parts)
            except Exception as exc2:
                last_exc = exc2
                if not _is_quota_error(exc2):
                    raise
        raise last_exc


def format_user_error(exc: Exception, language: str = "uz") -> str:
    """Texnik xatoni foydalanuvchi tilidagi tushunarli matnga aylantiradi."""
    if _is_quota_error(exc):
        # retry_delay (sekundlar) ni topishga harakat qilamiz
        import re
        m = re.search(r"retry[_ ]delay[^0-9]*(\d+)", str(exc), re.IGNORECASE)
        wait = m.group(1) if m else None
        wait_text = f" Taxminan {wait} soniyadan keyin qayta urinib ko'ring." if wait else ""
        messages = {
            "uz": "AI xizmati hozirda band — kunlik bepul so'rovlar chegarasiga yetildi.{wait}\n\nIltimos, biroz kutib qayta urinib ko'ring.",
            "ru": "Сервис ИИ сейчас перегружен — достигнут дневной лимит бесплатных запросов.{wait}\n\nПопробуйте позже.",
            "en": "The AI service is busy — the daily free request limit has been reached.{wait}\n\nPlease try again later.",
        }
        return messages.get(language, messages["uz"]).format(wait=wait_text)
    messages = {
        "uz": f"Xatolik yuz berdi: {exc}",
        "ru": f"Произошла ошибка: {exc}",
        "en": f"An error occurred: {exc}",
    }
    return messages.get(language, messages["uz"])

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

TELEGRAM_MAX_MESSAGE = 4000
MAX_DOC_TEXT_CHARS = 30000

# Qo'llab-quvvatlanadigan MIME turlari va kengaytmalar
PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
XLSX_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}
IMAGE_MIMES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/webp",
    "image/heic",
    "image/heif",
}
SUPPORTED_EXTS = (
    ".pdf", ".docx", ".xlsx",
    ".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif",
)


# ---------------------------------------------------------------------------
# Matn ajratuvchilar
# ---------------------------------------------------------------------------
MAX_OCR_PAGES = 10
OCR_DPI = 200


def extract_pdf_text(file_bytes: bytes) -> str:
    parts: list[str] = []
    with pymupdf.open(stream=io.BytesIO(file_bytes), filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts).strip()


def render_pdf_pages_as_images(file_bytes: bytes) -> list[dict]:
    """Skaner qilingan PDF sahifalarini PNG rasmga o'giradi (Gemini vision uchun)."""
    images: list[dict] = []
    with pymupdf.open(stream=io.BytesIO(file_bytes), filetype="pdf") as doc:
        for i, page in enumerate(doc):
            if i >= MAX_OCR_PAGES:
                break
            pix = page.get_pixmap(dpi=OCR_DPI)
            images.append(
                {"mime_type": "image/png", "data": pix.tobytes("png")}
            )
    return images


def extract_docx_text(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Paragraflar
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Jadvallar
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_xlsx_text(file_bytes: bytes) -> str:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Varaq: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [
                "" if v is None else str(v).strip()
                for v in row
            ]
            if any(c for c in cells):
                parts.append(" | ".join(cells))

    wb.close()
    return "\n".join(parts).strip()


def extract_text(file_bytes: bytes, mime: str, filename: str) -> tuple[str, str]:
    """Hujjat turiga qarab matnni ajratadi. (matn, hujjat_turi) qaytaradi."""
    name_lower = (filename or "").lower()

    if mime in PDF_MIMES or name_lower.endswith(".pdf"):
        return extract_pdf_text(file_bytes), "PDF"
    if mime in DOCX_MIMES or name_lower.endswith(".docx"):
        return extract_docx_text(file_bytes), "Word (.docx)"
    if mime in XLSX_MIMES or name_lower.endswith(".xlsx"):
        return extract_xlsx_text(file_bytes), "Excel (.xlsx)"

    raise ValueError("UNSUPPORTED")


def extract_inputs_for_ai(
    file_bytes: bytes, mime: str, filename: str
) -> tuple[str | None, list[dict] | None, str]:
    """
    Hujjat / rasmni AI uchun tayyorlaydi.
    Qaytaradi: (matn, rasmlar, hujjat_turi). Aynan biri to'ldirilgan bo'ladi.
    """
    name_lower = (filename or "").lower()
    image_exts = (".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif")

    if mime in IMAGE_MIMES or name_lower.endswith(image_exts):
        image_mime = mime if mime in IMAGE_MIMES else "image/jpeg"
        return None, [{"mime_type": image_mime, "data": file_bytes}], "Rasm (OCR)"

    if mime in PDF_MIMES or name_lower.endswith(".pdf"):
        text = extract_pdf_text(file_bytes)
        if text:
            return text, None, "PDF"
        images = render_pdf_pages_as_images(file_bytes)
        return None, images, "PDF (skaner / OCR)"

    if mime in DOCX_MIMES or name_lower.endswith(".docx"):
        return extract_docx_text(file_bytes), None, "Word (.docx)"

    if mime in XLSX_MIMES or name_lower.endswith(".xlsx"):
        return extract_xlsx_text(file_bytes), None, "Excel (.xlsx)"

    raise ValueError("UNSUPPORTED")


# ---------------------------------------------------------------------------
# Yordamchi funksiyalar
# ---------------------------------------------------------------------------
def chunk_text(text: str, size: int = TELEGRAM_MAX_MESSAGE) -> list[str]:
    if len(text) <= size:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            newline = text.rfind("\n", start, end)
            if newline > start + size // 2:
                end = newline
        chunks.append(text[start:end])
        start = end
    return chunks


def normalize_markdown(text: str) -> str:
    """AI ba'zan ** (CommonMark) qaytaradi — Telegram uchun bitta * ga aylantiramiz."""
    import re
    # **bold** -> *bold*
    text = re.sub(r"\*\*(.+?)\*\*", r"*\1*", text, flags=re.DOTALL)
    # __italic__ -> _italic_
    text = re.sub(r"__(.+?)__", r"_\1_", text, flags=re.DOTALL)
    return text


async def send_audit_message(message_obj, text: str, *, edit: bool = False) -> None:
    """Markdown bilan yuborishga harakat qiladi; xato bo'lsa oddiy matn sifatida yuboradi."""
    text = normalize_markdown(text)
    try:
        if edit:
            await message_obj.edit_text(text, parse_mode="Markdown")
        else:
            await message_obj.reply_text(text, parse_mode="Markdown")
    except Exception as exc:
        logger.warning("Markdown rendering failed (%s); plain text yuborilmoqda", exc)
        if edit:
            await message_obj.edit_text(text)
        else:
            await message_obj.reply_text(text)


def result_format_keyboard(context: ContextTypes.DEFAULT_TYPE) -> InlineKeyboardMarkup:
    language = context.user_data.get("language", "uz")
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(
                    TEXTS["as_text"][language], callback_data="audit_format_text"
                ),
                InlineKeyboardButton(
                    TEXTS["as_file"][language], callback_data="audit_format_file"
                ),
            ]
        ]
    )


async def offer_result_format(
    update: Update, context: ContextTypes.DEFAULT_TYPE, text: str
) -> None:
    context.user_data["pending_audit_result"] = text
    await update.message.reply_text(
        tr(context, "choose_result"),
        reply_markup=result_format_keyboard(context),
    )


async def handle_format_choice(
    update: Update, context: ContextTypes.DEFAULT_TYPE
) -> None:
    query = update.callback_query
    if query is None:
        return
    await query.answer()

    result_text = context.user_data.pop("pending_audit_result", None)
    if not result_text:
        await query.edit_message_text(tr(context, "result_missing"))
        return

    if query.data == "audit_format_text":
        await query.edit_message_text(tr(context, "result_sending_text"))
        for part in chunk_text(result_text):
            await context.bot.send_message(
                chat_id=query.from_user.id,
                text=part,
            )
        return

    await query.edit_message_text(tr(context, "result_sending_file"))
    file_buffer = io.BytesIO(result_text.encode("utf-8"))
    file_buffer.name = "audit_xulosasi.txt"
    await context.bot.send_document(
        chat_id=query.from_user.id,
        document=file_buffer,
        filename="audit_xulosasi.txt",
        caption="Audit xulosasi",
    )


def subscription_keyboard(context: ContextTypes.DEFAULT_TYPE) -> InlineKeyboardMarkup:
    language = context.user_data.get("language", "uz")
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(
                    TEXTS["join_channel"][language],
                    url=f"https://t.me/{CHANNEL_USERNAME.lstrip('@')}",
                )
            ],
            [
                InlineKeyboardButton(
                    TEXTS["check_subscription"][language], callback_data="check_channel_subscription"
                )
            ],
        ]
    )


def language_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton("O'zbekcha", callback_data="language_uz"),
                InlineKeyboardButton("Русский", callback_data="language_ru"),
                InlineKeyboardButton("English", callback_data="language_en"),
            ]
        ]
    )


async def is_channel_subscriber(
    update: Update, context: ContextTypes.DEFAULT_TYPE
) -> bool:
    user = update.effective_user
    if user is None:
        return False
    try:
        member = await context.bot.get_chat_member(
            chat_id=CHANNEL_USERNAME,
            user_id=user.id,
        )
        return member.status in {"creator", "administrator", "member"} or (
            member.status == "restricted" and bool(getattr(member, "is_member", False))
        )
    except Exception as exc:
        logger.warning("Kanal obunasini tekshirib bo'lmadi: %s", exc)
        return False


async def require_channel_subscription(
    update: Update, context: ContextTypes.DEFAULT_TYPE
) -> bool:
    if await is_channel_subscriber(update, context):
        return True
    await update.message.reply_text(
        tr(context, "must_subscribe"),
        reply_markup=subscription_keyboard(context),
    )
    return False


async def check_channel_subscription(
    update: Update, context: ContextTypes.DEFAULT_TYPE
) -> None:
    query = update.callback_query
    if query is None:
        return
    if await is_channel_subscriber(update, context):
        await query.answer(TEXTS["subscription_ok"][context.user_data.get("language", "uz")])
        await query.edit_message_text(
            tr(context, "subscription_ok")
        )
    else:
        await query.answer(tr(context, "not_subscribed"), show_alert=True)


async def choose_language(
    update: Update, context: ContextTypes.DEFAULT_TYPE
) -> None:
    query = update.callback_query
    if query is None or not query.data:
        return
    language = query.data.removeprefix("language_")
    if language not in LANGUAGES:
        await query.answer()
        return
    context.user_data["language"] = language
    await query.answer()
    if not await is_channel_subscriber(update, context):
        await query.edit_message_text(
            tr(context, "must_subscribe"),
            reply_markup=subscription_keyboard(context),
        )
        return
    await query.edit_message_text(tr(context, "welcome"))


SYSTEM_PROMPT = (
    "Sen professional o'zbek buxgalterisan. Berilgan hujjatdagi raqamlar, "
    "sanalar va tafovutlarni juda aniq tahlil qilib, o'zbek tilida xulosa ber."
)

LANGUAGE_INSTRUCTIONS = {
    "uz": "Javobni o'zbek tilida ber.",
    "ru": "Дай ответ на русском языке.",
    "en": "Give the answer in English.",
}

LANGUAGE_NAMES = {
    "uz": "o'zbek",
    "ru": "рус",
    "en": "English",
}


BUXGALTERIYA_SYSTEM_PROMPT_TEMPLATE = """
Sen tajribali buxgalter-maslahatchisan. Faqat buxgalteriya, soliq, moliyaviy
hisobot, audit va shu kabi professional mavzularga oid savollarga javob ber.

TIL QOIDASI:
- Har doim {preferred_lang} tilida javob ber — savol qaysi tilda yozilganidan
  qat'iy nazar.

QOIDALAR:
- Agar savolda aniq mamlakat yoki standart ko'rsatilmagan bo'lsa, kerak bo'lsa
  O'zbekiston amaliyoti va xalqaro IFRS amaliyotini qisqacha qiyosla.
- Agar savolda aniq mamlakat yoki standart ko'rsatilgan bo'lsa (masalan IFRS,
  AQSh yoki Rossiya), faqat o'shanga tayan.
- Savol buxgalteriya, soliq yoki moliya bilan bog'liq bo'lmasa, {preferred_lang}
  tilida muloyimlik bilan bu mavzuda yordam bera olmasligingni ayt.
- Javoblar qisqa, aniq va amaliy bo'lsin — uzun ma'ruza emas.
- Har bir javobni quyidagi tartibda yoz:
  1. Savolning aniq ma'nosi yoki ta'rifi — foydalanuvchi nimani so'rayotganini
     bir-ikki jumlada tushuntir.
  2. Oddiy hayotiy yoki buxgalteriya misoli keltir.
  3. Natijani sodda, hammaga tushunarli tilda izohla.
- Shu uch qismni tanlangan tilda yoz. Zarur bo'lsa, oxirida qisqa amaliy
  tavsiya ber, lekin javobni ortiqcha cho'zma.
- Aniq stavka yoki raqam so'ralsa, joriy ma'lumotni rasmiy manbadan tekshirish
  kerakligini eslat, chunki qonunlar tez o'zgaradi.
- Murakkab savollarda professional buxgalter yoki auditor bilan maslahatlashish
  tavsiya etilishini qo'sh.
"""


def language_instruction(language: str) -> str:
    return LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["uz"])


_AUDIT_FORMAT_RULES = (
    "Sen auditorlik botisan. Faqat eng muhim tafovutlarni QISQA formatda "
    "ko'rsat. Gaplarni cho'zma. Telegram uchun chiroyli dizaynda yoz.\n\n"
    "Agar hujjatda bank ko'chirmasi bo'lsa — uni HAQIQAT deb hisobla "
    "(\"Bank/Bizda\" deb belgila). Aks holda \"Bizda\" deb yoz.\n\n"
    "Javob formati AYNAN quyidagicha bo'lsin (boshqa hech narsa qo'shma):\n\n"
    "📍 *TAFOVUTLAR:*\n"
    "`------------------------------`\n"
    "🗓 *15.04.2026 | №44 To'lov*\n"
    "❌ Hamkor: 12 mln (Kredit)\n"
    "✅ Bank/Bizda: 15,000,000 (Debit)\n"
    "⚠️ _Farq: 3,000,000 so'm_\n"
    "\n"
    "🗓 *20.04.2026 | №105 Xizmat*\n"
    "❌ Hamkor: 22.04 da yozgan\n"
    "✅ Bizda: 20.04 da yozilgan\n"
    "`------------------------------`\n"
    "\n"
    "⚠️ *XULOSA:* Hamkor-Trade Debit/Kreditda adashgan.\n"
    "\n"
    "💰 *HAQIQIY QOLDIQ:* 15,000,000 so'm (Bizning foydaga).\n\n"
    "QOIDALAR:\n"
    "• Har bir tafovutni alohida 🗓 kartochka qilib yoz (sana | hujjat raqami "
    "va turi).\n"
    "• ❌ — kim adashgan tomon, ✅ — to'g'ri tomon.\n"
    "• Qalin matn uchun BITTA yulduzcha (*matn*), kursiv uchun (_matn_), "
    "ajratuvchi chiziq uchun teskari apostrof (`---`).\n"
    "• Ikkita yulduzcha (**) ISHLATMA — Telegram tushunmaydi.\n"
    "• Kirish so'zi, hujjat tavsifi, \"hurmatli\" kabi gaplarni YOZMA — "
    "to'g'ridan-to'g'ri 📍 dan boshla.\n"
    "• Agar tafovut yo'q bo'lsa, faqat: \"✅ Hujjatda tafovut aniqlanmadi.\""
)


def build_audit_prompt(
    doc_text: str, doc_type: str, language: str = "uz"
) -> str:
    truncated = doc_text[:MAX_DOC_TEXT_CHARS]
    return (
        f"{_AUDIT_FORMAT_RULES}\n{language_instruction(language)}\n\n"
        f"Hujjat turi: {doc_type}.\n\n"
        "=== HUJJAT MATNI ===\n"
        f"{truncated}\n"
        "=== HUJJAT TUGADI ==="
    )


def build_ocr_prompt(
    doc_type: str, page_count: int, language: str = "uz"
) -> str:
    return (
        f"{_AUDIT_FORMAT_RULES}\n{language_instruction(language)}\n\n"
        f"Quyidagi {page_count} ta rasm — skaner qilingan hujjat sahifalari "
        f"({doc_type}). Rasmlardagi raqam, sana, jadval va imzolarni diqqat "
        "bilan o'qib, yuqoridagi formatda javob ber. Agar biror raqam "
        "noaniq bo'lsa, qavs ichida [noaniq] deb belgila."
    )


def build_compare_prompt(
    text1: str, name1: str, text2: str, name2: str, language: str = "uz"
) -> str:
    return f"""
Sen professional auditorsan. Quyida ikkita moliyaviy hisobot berilgan — ularni
solishtirib, FAQAT aniq tafovutlarni top. Mos keladigan, farqsiz tranzaksiyalarni
sanab o'tirma, ular haqida umuman yozma — faqat MUAMMOLARNI ko'rsat.
{language_instruction(language)}

QOIDALAR:
- Har ikki fayldagi yozuvlarni sana, hujjat raqami (agar bo'lsa) va summa bo'yicha moslashtir.
- Faqat quyidagi holatlarni tafovut deb hisobla va ro'yxatla:
  1. Summasi ikki faylda har xil bo'lsa — aniq raqam bilan farqni ko'rsat.
  2. Bitta faylda bor, ikkinchisida yo'q yozuv.
  3. Debit/Kredit tomoni chalkashgan bo'lsa.
  4. Sana boshqa, lekin summa yoki hujjat raqami bir xil bo'lsa.
- Yaxshi mos kelgan yozuvlarni sanama va umuman aytma.
- G'ayrioddiy katta summa haqida ogohlantirma — faqat aniq raqamli
  nomuvofiqlikni ko'rsat.
- Ma'lumot yetarli bo'lmasa, summa yoki xulosani o'ylab topma.
- Agar tafovut umuman topilmasa, faqat: "Tafovut aniqlanmadi" deb yoz.

JAVOB FORMATI — QAT'IY SHUNGA AMAL QIL:

AUDIT XULOSASI
------------------------------
[Har bir tafovut uchun faqat muammoli yozuvlarni ko'rsat]

🗓 [sana] | №[hujjat raqami yoki nomi]
❌ {name1}: [summa] ([Debit/Kredit])
✅ {name2}: [summa] ([Debit/Kredit])
⚠️ Farq: [aniq raqam] so'm
------------------------------

[Agar bir faylda umuman yo'q bo'lsa:]
🗓 [sana] | №[hujjat raqami]
❌ {name1}: yo'q
✅ {name2}: [summa] ([Debit/Kredit])
------------------------------

XULOSA: [bir jumlada, ehtimoliy sabab]
HAQIQIY QOLDIQ FARQI: [umumiy farq raqami] so'm ([kimning foydasiga])

Telegram uchun qisqa, o'qishga qulay satrlardan foydalan. Qo'shimcha kirish
so'zi, mos tranzaksiyalar ro'yxati yoki umumiy maslahatlar yozma.

--- FAYL 1: {name1} ---
{text1[:15000]}

--- FAYL 2: {name2} ---
{text2[:15000]}
"""


# ---------------------------------------------------------------------------
# Telegram handlerlari
# ---------------------------------------------------------------------------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if "language" not in context.user_data:
        await update.message.reply_text(
            "Tilni tanlang / Выберите язык / Choose your language:",
            reply_markup=language_keyboard(),
        )
        return

    if not await is_channel_subscriber(update, context):
        await update.message.reply_text(
            tr(context, "must_subscribe"),
            reply_markup=subscription_keyboard(context),
        )
        return

    await update.message.reply_text(tr(context, "welcome"))


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    help_text = {
        "uz": "Foydalanish:\n1. PDF, Word, Excel hujjat yoki rasm yuboring.\n2. Men uni AI yordamida tahlil qilaman.\n3. Sizga audit xulosasini qaytaraman.\n\nIkkita hujjatni solishtirish uchun /compare buyrug'ini ishlating.\nBekor qilish: /cancel",
        "ru": "Использование:\n1. Отправьте PDF, Word, Excel-документ или изображение.\n2. Я проанализирую его с помощью ИИ.\n3. Вы получите аудиторское заключение.\n\nДля сравнения двух документов используйте /compare.\nОтмена: /cancel",
        "en": "How to use:\n1. Send a PDF, Word, Excel document, or image.\n2. I will analyze it with AI.\n3. You will receive an audit summary.\n\nUse /compare to compare two documents.\nCancel: /cancel",
    }
    await update.message.reply_text(help_text.get(context.user_data.get("language", "uz"), help_text["uz"]))


async def compare_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not await require_channel_subscription(update, context):
        return
    context.user_data["compare_mode"] = True
    context.user_data.pop("compare_first", None)
    await update.message.reply_text(tr(context, "compare_started"))


async def cancel_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    was_active = context.user_data.pop("compare_mode", False)
    context.user_data.pop("compare_first", None)
    if was_active:
        await update.message.reply_text(tr(context, "cancelled"))
    else:
        await update.message.reply_text(tr(context, "nothing_to_cancel"))


async def handle_text_question(
    update: Update, context: ContextTypes.DEFAULT_TYPE
) -> None:
    if not await require_channel_subscription(update, context):
        return

    language = context.user_data.get("language", "uz")
    preferred_lang = LANGUAGE_NAMES.get(language, LANGUAGE_NAMES["uz"])
    thinking_message = {
        "uz": "🤔 O'ylab ko'ryapman...",
        "ru": "🤔 Думаю...",
        "en": "🤔 Thinking...",
    }
    await update.message.reply_text(thinking_message.get(language, thinking_message["uz"]))

    question = update.message.text.strip()
    system_prompt = BUXGALTERIYA_SYSTEM_PROMPT_TEMPLATE.format(
        preferred_lang=preferred_lang
    )
    try:
        response = generate_with_fallback(
            f"{system_prompt}\n\nSavol:\n{question}"
        )
        answer = (response.text or "").strip()
        if not answer:
            await update.message.reply_text(tr(context, "no_ai_response"))
            return
        for part in chunk_text(answer):
            await update.message.reply_text(part)
    except Exception as exc:
        logger.exception("Matnli savolga javob berishda xatolik")
        await update.message.reply_text(
            format_user_error(exc, language)
        )


async def _run_comparison(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    status_message,
    first: dict,
    second: dict,
) -> None:
    """Ikki hujjatni Gemini-ga solishtirish uchun yuboradi."""
    first_name = first.get("name", first["doc_type"])
    second_name = second.get("name", second["doc_type"])
    prompt = build_compare_prompt(
        first["text"] or "",
        first_name,
        second["text"] or "",
        second_name,
        context.user_data.get("language", "uz"),
    )
    parts: list = [prompt, f"\n=== HUJJAT 1 ({first_name}) ===\n"]
    if first["text"]:
        parts.append(first["text"][:MAX_DOC_TEXT_CHARS])
    else:
        parts.extend(first["images"])
    parts.append("\n=== HUJJAT 1 TUGADI ===\n")
    parts.append(f"\n=== HUJJAT 2 ({second_name}) ===\n")
    if second["text"]:
        parts.append(second["text"][:MAX_DOC_TEXT_CHARS])
    else:
        parts.extend(second["images"])
    parts.append("\n=== HUJJAT 2 TUGADI ===\n")

    response = generate_with_fallback(parts)
    audit_text = (response.text or "").strip()

    if not audit_text:
        await status_message.edit_text(
            tr(context, "no_ai_response")
        )
        return

    header = (
        f"SOLISHTIRISH XULOSASI\n"
        f"({first['doc_type']} vs {second['doc_type']})\n\n"
    )
    await offer_result_format(update, context, header + audit_text)


async def _process_compare_file(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    file_bytes: bytes,
    mime: str,
    filename: str,
) -> None:
    """Solishtirish rejimida kelgan faylni qayta ishlaydi."""
    status_message = await update.message.reply_text(
        tr(context, "file_received")
    )
    try:
        text, images, doc_type = extract_inputs_for_ai(file_bytes, mime, filename)
    except ValueError:
        await status_message.edit_text(tr(context, "unsupported"))
        return

    if text is None and not images:
        await status_message.edit_text(tr(context, "no_data"))
        return

    first = context.user_data.get("compare_first")
    if first is None:
        context.user_data["compare_first"] = {
            "text": text,
            "images": images,
            "doc_type": doc_type,
            "name": filename or doc_type,
        }
        await status_message.edit_text(
            TEXTS["second_document"][context.user_data.get("language", "uz")].format(
                doc_type=doc_type
            )
        )
        return

    second = {
        "text": text,
        "images": images,
        "doc_type": doc_type,
        "name": filename or doc_type,
    }
    await status_message.edit_text(
        tr(context, "both_documents")
    )
    try:
        await _run_comparison(update, context, status_message, first, second)
    finally:
        context.user_data.pop("compare_mode", None)
        context.user_data.pop("compare_first", None)


async def _run_audit_on_images(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    status_message,
    images: list[dict],
    doc_type: str,
) -> None:
    """Rasmlar ro'yxatini Gemini vision ga yuborib audit xulosasini chiqaradi."""
    language = context.user_data.get("language", "uz")
    ocr_prompt = build_ocr_prompt(
        doc_type, page_count=len(images), language=language
    )
    response = generate_with_fallback([ocr_prompt, *images])
    audit_text = (response.text or "").strip()
    logger.info("Vision rejimida %d sahifa/rasm tahlil qilindi", len(images))

    if not audit_text:
        await status_message.edit_text(
            tr(context, "no_ai_response")
        )
        return

    header = f"AUDIT XULOSASI ({doc_type})\n\n"
    chunks = chunk_text(header + audit_text)
    await send_audit_message(status_message, chunks[0], edit=True)
    for extra in chunks[1:]:
        await send_audit_message(update.message, extra)


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    document = update.message.document
    if document is None:
        await update.message.reply_text(
            tr(context, "other_prompt")
        )
        return
    if not await require_channel_subscription(update, context):
        return

    filename = document.file_name or ""
    mime = document.mime_type or ""
    name_lower = filename.lower()

    is_image = mime in IMAGE_MIMES or name_lower.endswith(
        (".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif")
    )
    is_supported = (
        is_image
        or mime in PDF_MIMES
        or mime in DOCX_MIMES
        or mime in XLSX_MIMES
        or name_lower.endswith(SUPPORTED_EXTS)
    )
    if not is_supported:
        await update.message.reply_text(tr(context, "unsupported"))
        return

    # Solishtirish rejimi yoqilgan bo'lsa — alohida oqimga yuboramiz
    if context.user_data.get("compare_mode"):
        try:
            file = await context.bot.get_file(document.file_id)
            file_bytes = bytes(await file.download_as_bytearray())
            await _process_compare_file(
                update, context, file_bytes, mime, filename
            )
        except Exception as exc:
            logger.exception("Solishtirish rejimida xatolik")
            await update.message.reply_text(
                format_user_error(exc, context.user_data.get("language", "uz"))
            )
        return

    status_message = await update.message.reply_text(
        tr(context, "analysis_wait")
    )

    try:
        file = await context.bot.get_file(document.file_id)
        file_bytes = bytes(await file.download_as_bytearray())

        # Rasm sifatida yuborilgan fayl — to'g'ridan-to'g'ri vision OCR
        if is_image:
            image_mime = mime if mime in IMAGE_MIMES else "image/jpeg"
            images = [{"mime_type": image_mime, "data": file_bytes}]
            await _run_audit_on_images(
                update, context, status_message, images, "Rasm (OCR)"
            )
            return

        try:
            doc_text, doc_type = extract_text(file_bytes, mime, filename)
        except ValueError:
            await status_message.edit_text(tr(context, "unsupported"))
            return

        used_ocr = False
        if not doc_text and doc_type == "PDF":
            await status_message.edit_text(tr(context, "ocr_wait"))
            try:
                images = render_pdf_pages_as_images(file_bytes)
            except Exception as exc:
                logger.exception("PDF sahifalarini rasmga o'girishda xatolik")
                await status_message.edit_text(
                    f"PDF sahifalarini rasmga o'girib bo'lmadi: {exc}"
                )
                return

            if not images:
                await status_message.edit_text(tr(context, "no_data"))
                return

            used_ocr = True
            ocr_prompt = build_ocr_prompt(
                doc_type,
                page_count=len(images),
                language=context.user_data.get("language", "uz"),
            )
            response = generate_with_fallback([ocr_prompt, *images])
            doc_type = "PDF (skaner / OCR)"
        elif not doc_text:
            await status_message.edit_text(tr(context, "empty_document"))
            return
        else:
            logger.info(
                "%s '%s' qabul qilindi, %d ta belgi olindi",
                doc_type,
                filename,
                len(doc_text),
            )
            prompt = build_audit_prompt(
                doc_text,
                doc_type,
                language=context.user_data.get("language", "uz"),
            )
            response = generate_with_fallback(prompt)

        audit_text = (response.text or "").strip()
        if used_ocr:
            logger.info("OCR rejimida %d sahifa tahlil qilindi", len(images))

        if not audit_text:
            await status_message.edit_text(tr(context, "no_ai_response"))
            return

        header = f"AUDIT XULOSASI ({doc_type})\n\n"
        await offer_result_format(update, context, header + audit_text)

    except Exception as exc:
        logger.exception("Hujjatni qayta ishlashda xatolik")
        await status_message.edit_text(
            format_user_error(exc, context.user_data.get("language", "uz"))
        )


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not await require_channel_subscription(update, context):
        return
    photos = update.message.photo
    if not photos:
        await update.message.reply_text(tr(context, "other_prompt"))
        return

    # Eng katta o'lchamdagi versiyani tanlaymiz
    photo = photos[-1]

    # Solishtirish rejimi yoqilgan bo'lsa — alohida oqimga yuboramiz
    if context.user_data.get("compare_mode"):
        try:
            file = await context.bot.get_file(photo.file_id)
            file_bytes = bytes(await file.download_as_bytearray())
            await _process_compare_file(
                update, context, file_bytes, "image/jpeg", "photo.jpg"
            )
        except Exception as exc:
            logger.exception("Solishtirish rejimida (rasm) xatolik")
            await update.message.reply_text(
                format_user_error(exc, context.user_data.get("language", "uz"))
            )
        return

    status_message = await update.message.reply_text(
        tr(context, "image_wait")
    )

    try:
        file = await context.bot.get_file(photo.file_id)
        file_bytes = bytes(await file.download_as_bytearray())
        images = [{"mime_type": "image/jpeg", "data": file_bytes}]
        await _run_audit_on_images(
            update, context, status_message, images, "Rasm (OCR)"
        )
    except Exception as exc:
        logger.exception("Rasmni qayta ishlashda xatolik")
        await status_message.edit_text(
            format_user_error(exc, context.user_data.get("language", "uz"))
        )


async def handle_other(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(tr(context, "other_prompt"))


# ---------------------------------------------------------------------------
# Bot ishga tushirish
# ---------------------------------------------------------------------------
def main() -> None:
    app = ApplicationBuilder().token(TELEGRAM_BOT_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("compare", compare_command))
    app.add_handler(CommandHandler("cancel", cancel_command))
    app.add_handler(
        CallbackQueryHandler(handle_format_choice, pattern=r"^audit_format_(text|file)$")
    )
    app.add_handler(
        CallbackQueryHandler(choose_language, pattern=r"^language_(uz|ru|en)$")
    )
    app.add_handler(
        CallbackQueryHandler(
            check_channel_subscription, pattern=r"^check_channel_subscription$"
        )
    )
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(
        MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text_question)
    )
    app.add_handler(
        MessageHandler(
            filters.ALL
            & ~filters.COMMAND
            & ~filters.Document.ALL
            & ~filters.PHOTO,
            handle_other,
        )
    )

    logger.info("ACCO AI bot ishga tushdi (PDF / Word / Excel / Rasm)...")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
